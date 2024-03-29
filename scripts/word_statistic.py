# -*- encoding: utf-8 -*-
"""
@File Name      :   word_statistic.py    
@Create Time    :   2022/8/31 15:13
@Description    :   
@Version        :   
@License        :   MIT
@Author         :   diklios
@Contact Email  :   diklios5768@gmail.com
@Github         :   https://github.com/diklios5768
@Blog           :   
@Motto          :   All our science, measured against reality, is primitive and childlike - and yet it is the most precious thing we have.
"""
__auth__ = 'diklios'

import os
from collections import defaultdict

import PyPDF2
import click
import docx
from tqdm import tqdm


def merge_dict(x: defaultdict, y: dict):
    for k, v in y.items():
        x[k] += v
    return x


def is_file_type(file_path: str, file_type: str):
    return os.path.isfile(file_path) and '~$' not in file_path and file_path.endswith(file_type)


def is_docx_file(file_path: str):
    return is_file_type(file_path, '.docx')


def is_doc_file(file_path: str):
    return is_file_type(file_path, '.doc')


def is_pdf_file(file_path: str):
    return is_file_type(file_path, '.pdf')


def is_txt_file(file_path: str):
    return is_file_type(file_path, '.txt')


def is_need_file(file_path: str):
    return is_docx_file(file_path) or is_pdf_file(file_path)


def is_word_character(character: str):
    return 'a' <= character <= 'z' or 'A' <= character <= 'Z' or character == '-'


class WordStatistic:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.hash = defaultdict(int)

    def sort_hash_by_word(self):
        #  sort by word
        self.hash = dict(sorted(self.hash.items(), key=lambda x: x[0]))
        self.hash.pop('')
        return self.hash

    def count_word(self, text: str):
        word = ''
        for character in text:
            if is_word_character(character):
                word += character.lower()
            else:
                self.hash[word] += 1
                word = ''

    def statistic_docx(self):
        file = docx.Document(self.file_path)
        for page_num in range(len(file.paragraphs)):
            self.count_word(file.paragraphs[page_num].text)
        return self.sort_hash_by_word()

    def statistic_doc(self):
        pass

    def statistic_pdf(self):
        file = PyPDF2.PdfFileReader(open(self.file_path, 'rb'))
        for page_num in range(file.numPages):
            self.count_word(file.getPage(page_num).extractText().strip().replace('\n', ' '))
        return self.sort_hash_by_word()

    def statistic_txt(self):
        pass

    def statistic(self):
        if is_docx_file(self.file_path):
            return self.statistic_docx()
        elif is_pdf_file(self.file_path):
            return self.statistic_pdf()
        else:
            raise Exception(f'{self.file_path} file type not supported!')

    def __call__(self):
        return self.statistic()


@click.command()
@click.argument('dir_path', type=click.Path(exists=True))
def main(dir_path: str):
    statistic_all = defaultdict(int)
    file_paths = [os.path.join(dir_path, file_name) for file_name in os.listdir(dir_path)]
    for file_path in tqdm([file_path for file_path in file_paths if is_need_file(file_path)]):
        try:
            statistic_one = WordStatistic(file_path).statistic()
        except Exception as e:
            print(e)
        else:
            statistic_all = merge_dict(statistic_all, statistic_one)
    # sort by frequency
    statistic_all = dict(sorted(statistic_all.items(), key=lambda x: x[1], reverse=True))
    # generate file
    with open(os.path.join(dir_path, 'word_statistic.txt'), 'w') as f:
        f.write('\n'.join([f'{word}:{frequency}' for word, frequency in statistic_all.items()]))
    print('statistic success')


if __name__ == '__main__':
    """
    contributor:https://github.com/chenyuzhe97
    usage:python word_statistic.py dir_path
    
    todo:doc,txt file support
    """
    main()
